from docx import Document
from langchain.chains import LLMChain
from langchain.prompts import PromptTemplate
from langchain_community.llms import Ollama
import os

# Ollama LLM 초기화
llm = Ollama(model="mistral:latest")

# 프롬프트 템플릿 설정
prompt = PromptTemplate(
    input_variables=["claim_text"],
    template="다음 내용을 법적 서류에 맞는 단어와 말투로 수정해줘: {claim_text}"
)

# '청구 취지' 또는 '청구 원인'을 찾는 함수
def find_section(paragraphs, keywords):
    for i, paragraph in enumerate(paragraphs):
        if any(keyword in paragraph.text for keyword in keywords):
            return i
    return None

def fix_doc(file_path: str):
    # 파일을 Document 객체로 로드
    doc = Document(file_path)
    
    # 문서에서 모든 텍스트 추출
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)

    # 모든 텍스트를 하나의 문자열로 결합
    combined_text = "\n".join(full_text)

    # 청구 취지와 청구 원인 부분을 찾아 수정
    claim_purpose_keywords = ["청구취지", "청구  취지", "청  구  취  지"]
    claim_reason_keywords = ["청구원인", "청구  원인", "청  구  원  인"]

    # 각 부분에 대해 텍스트를 수정
    def modify_section(text):
        llm_chain = LLMChain(llm=llm, prompt=prompt)
        return llm_chain.run(claim_text=text)

    # 텍스트 수정 프로세스 (청구 취지와 청구 원인 찾기 및 수정)
    claim_purpose_index = find_section(doc.paragraphs, claim_purpose_keywords)
    if claim_purpose_index is not None:
        original_text = doc.paragraphs[claim_purpose_index + 1].text
        modified_text = modify_section(original_text)
        doc.paragraphs[claim_purpose_index + 1].text = modified_text

    claim_reason_index = find_section(doc.paragraphs, claim_reason_keywords)
    if claim_reason_index is not None:
        original_text = doc.paragraphs[claim_reason_index + 1].text
        modified_text = modify_section(original_text)
        doc.paragraphs[claim_reason_index + 1].text = modified_text

    # 수정된 문서를 저장
    modified_file_path = "./output/fixed_document.docx"
    doc.save(modified_file_path)

    # 파일 경로를 반환
    return modified_file_path
